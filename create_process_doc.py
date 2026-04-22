from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── colour palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1F, 0x38, 0x64)
BLUE       = RGBColor(0x2E, 0x75, 0xB6)
LIGHTBLUE  = RGBColor(0xDD, 0xEB, 0xF7)
GREEN      = RGBColor(0x37, 0x86, 0x44)
LIGHTGREEN = RGBColor(0xE2, 0xEF, 0xDA)
ORANGE     = RGBColor(0xED, 0x7D, 0x31)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY       = RGBColor(0x40, 0x40, 0x40)
SILVER     = RGBColor(0xF2, 0xF2, 0xF2)

doc = Document()

# ── page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


# ── helpers ──────────────────────────────────────────────────────────────────
def hex_to_rgb_str(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"

def set_cell_bg(cell, hex_str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "6"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "BFBFBF"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(text, level=1, color=NAVY, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt({1: 20, 2: 15, 3: 12, 4: 11}.get(level, 11))
    if level == 1:
        # underline rule
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E75B6")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_body(text, indent=0, space_after=4, color=GREY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Inches(indent * 0.25)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = color
    return p

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + "  ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = BLUE
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GREY

def add_step_table(rows_data, col_widths=(1.0, 1.4, 3.9)):
    """rows_data = list of (step_no, label, description)"""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    # header
    hdr = tbl.rows[0].cells
    for i, (txt, w) in enumerate(zip(["Step", "Component", "What Happens"], col_widths)):
        hdr[i].width = Inches(w)
        set_cell_bg(hdr[i], "1F3864")
        p = hdr[i].paragraphs[0]
        p.clear()
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for idx, (step, label, desc) in enumerate(rows_data):
        row = tbl.add_row().cells
        bg = "F2F8FF" if idx % 2 == 0 else "FFFFFF"
        set_cell_bg(row[0], bg)
        set_cell_bg(row[1], bg)
        set_cell_bg(row[2], bg)

        # step number
        p0 = row[0].paragraphs[0]; p0.clear()
        r = p0.add_run(str(step))
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLUE
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # component label
        p1 = row[1].paragraphs[0]; p1.clear()
        r = p1.add_run(label)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY

        # description
        p2 = row[2].paragraphs[0]; p2.clear()
        r = p2.add_run(desc)
        r.font.size = Pt(10); r.font.color.rgb = GREY

        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)

    doc.add_paragraph()   # spacer


def add_two_col_table(rows_data, header=("Term / Component", "Description"),
                      col_widths=(1.8, 4.5)):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, (txt, w) in enumerate(zip(header, col_widths)):
        hdr[i].width = Inches(w)
        set_cell_bg(hdr[i], "2E75B6")
        p = hdr[i].paragraphs[0]; p.clear()
        run = p.add_run(txt)
        run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, (term, desc) in enumerate(rows_data):
        row = tbl.add_row().cells
        bg = "F2F2F2" if idx % 2 == 0 else "FFFFFF"
        set_cell_bg(row[0], bg); set_cell_bg(row[1], bg)
        p0 = row[0].paragraphs[0]; p0.clear()
        r = p0.add_run(term); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY
        p1 = row[1].paragraphs[0]; p1.clear()
        r = p1.add_run(desc); r.font.size = Pt(10); r.font.color.rgb = GREY
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)

    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
#  COVER / TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(48)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI-DA-AGENTS")
r.bold = True; r.font.size = Pt(32); r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Process & Architecture Document")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Retail Sales Analytics Chatbot  ·  Version 1.0  ·  April 2026")
r.font.size = Pt(11); r.font.color.rgb = GREY; r.italic = True

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1.  Project Overview", level=1)
add_body(
    "AI-DA-AGENTS is a multi-agent conversational analytics assistant that allows retail "
    "business users to ask questions in plain English and receive accurate, data-driven "
    "answers backed by live SQL queries against a Microsoft Fabric SQL endpoint.  "
    "The system handles everything from intent classification and entity resolution through "
    "to SQL generation, execution, visualisation, and natural-language summarisation — "
    "all within a single conversational interface."
)

add_heading("Key Capabilities", level=3, space_before=10)
bullets = [
    ("Natural-Language to SQL",   "Converts plain-English business questions into safe, read-only T-SQL."),
    ("Multi-Agent Architecture",  "Specialised LLM agents for classification, clarification, SQL writing, debugging, summarisation, and empty-result handling."),
    ("Entity Resolution",         "Fuzzy + embedding-based matching for STATE, CITY, STORE NAME, CATEGORY, and SUBCLASS values."),
    ("Semantic SQL Cache",        "Exact and cosine-similarity based cache avoids redundant LLM calls for repeated questions."),
    ("Conversation Memory",       "Persistent, thread-aware memory keeps context across turns and sessions."),
    ("Smart Visualisation",       "Data-driven chart selection (bar, line, pie, HTML table) — charts only when data warrants them."),
    ("Clarification & Recovery",  "Asks follow-up questions for vague queries; retries SQL automatically on database errors; explains empty results."),
]
for term, desc in bullets:
    add_bullet(desc, bold_prefix=term)


# ══════════════════════════════════════════════════════════════════════════════
#  2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2.  System Architecture", level=1)
add_body(
    "The system is organised around a central Orchestrator that coordinates a pipeline of "
    "specialised agents, utility services, and data stores.  All external LLM calls are "
    "routed through a unified LiteLLM wrapper, while database access is isolated behind a "
    "read-only execution layer."
)

add_heading("Core Layers", level=2, space_before=10, space_after=4)
add_two_col_table([
    ("Presentation Layer",   "Streamlit web app (app.py) — renders chat turns, SQL, charts, and table previews."),
    ("Orchestration Layer",  "ChatOrchestrator (orchestrator.py) — coordinates all agents and services for every user message."),
    ("Agent Layer",          "Six LLM-powered agents: Classifier, Clarifier, SQL Writer, SQL Debugger, Summariser, Empty Result Handler."),
    ("Entity Layer",         "Five column-specific entity resolvers (STATE, CITY, STORE_NAME, CATEGORY, SUBCLASS) using embeddings + fuzzy matching."),
    ("Cache Layer",          "SQLite-backed SQL query cache with exact-match and semantic (cosine) lookup."),
    ("Data Layer",           "Microsoft Fabric SQL endpoint — single analytics table: prd.FACT_SALES_AI."),
    ("Memory Layer",         "JSON-file backed conversation memory with multi-thread support."),
    ("Schema Layer",         "SchemaCache — retrieves and caches the live table schema from the database."),
], header=("Layer", "Responsibility"))


# ══════════════════════════════════════════════════════════════════════════════
#  3. END-TO-END MESSAGE FLOW
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3.  End-to-End Message Flow", level=1)
add_body(
    "Every user message passes through the following sequential pipeline.  "
    "Steps 6 and 10 run in parallel to minimise response latency."
)

add_step_table([
    (1,  "Input Validation",       "Strip whitespace; return prompt if empty."),
    (2,  "Intent Classification",  "Classifier Agent labels the message as 'normal_chat' or 'business_question' using conversation history + LLM."),
    (3,  "Normal Chat Branch",     "If normal_chat: Chat Agent responds conversationally; turn is stored in memory; flow ends."),
    (4,  "Schema & Context Load",  "SchemaCache returns the live table schema; BusinessContextStore provides column definitions and KPI rules."),
    (5,  "Clarification Check",    "Clarifier Agent analyses the question for missing critical context (time period, dimension, metric). If vague → asks follow-up and halts pipeline."),
    (6,  "Entity Resolution",      "Five resolvers run sequentially against the question. Matched entities (e.g. SUBCLASS=JOGGERS, STATE=MAHARASHTRA) are formatted as a hint block injected into the LLM prompt."),
    (7,  "SQL Cache Lookup",       "1) Exact match on normalised question + schema fingerprint + business fingerprint. 2) Semantic match via cosine similarity on question embedding. Returns cached SQL if found."),
    (8,  "SQL Generation",         "If cache misses: SQL Writer Agent generates a single read-only T-SQL SELECT using schema context, business context, entity hints, and KPI rules."),
    (9,  "SQL Guard",              "ensure_safe_readonly_sql() validates the SQL is SELECT/WITH-only before execution. Raises SqlGuardError if unsafe."),
    (10, "SQL Execution + Retry",  "execute_sql_query() runs the SQL. On DatabaseExecutionError or SqlGuardError: SQL Debugger Agent repairs the SQL and retries (up to sql_debug_max_retries)."),
    (11, "Empty Result Check",     "If row_count == 0: Empty Result Handler Agent generates a contextual follow-up citing the resolved entity attributes that were used."),
    (12, "Summarise + Visualise",  "Parallel threads: (a) Summariser Agent produces a natural-language answer; (b) build_visual_output() decides chart type and renders HTML chart or table."),
    (13, "Cache Update",           "Successful result is upserted into the SQL cache with question embedding."),
    (14, "Response Assembly",      "BotReply is built with answer text, SQL, row preview, chart path, and chart type."),
    (15, "Memory Storage",         "Turn (user message + assistant reply + route) is appended to the active conversation thread."),
])


# ══════════════════════════════════════════════════════════════════════════════
#  4. AGENT PROCESSES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4.  Agent Processes", level=1)
add_body("Each agent is a stateless function backed by a dedicated LLM call with fixed system instructions.")

agents = [
    ("4.1  Classifier Agent",
     "src/agents/classifier.py",
     "Determines whether the user's message requires a database query or is a general conversation.",
     [
         "Input: contextual prompt (conversation history + current message).",
         "Output: JSON  { label: 'business_question' | 'normal_chat' }.",
         "Falls back to 'normal_chat' on exception to prevent pipeline breakage.",
     ]),
    ("4.2  Clarifier Agent",
     "src/agents/clarifier.py",
     "Detects open-ended or incomplete business questions before any SQL is written.",
     [
         "Input: question with conversation history + business context.",
         "Checks: time period (most critical), dimension (region/store), metric.",
         "Output: JSON  { needs_clarification: bool, clarifying_question: str }.",
         "Reads history to avoid re-asking a question already answered in the thread.",
         "Falls back to needs_clarification=False on exception (never blocks unnecessarily).",
     ]),
    ("4.3  SQL Writer Agent",
     "src/agents/sql_writer.py",
     "Generates a single safe, read-only T-SQL SELECT statement for Microsoft Fabric SQL endpoint.",
     [
         "Only permitted table: prd.FACT_SALES_AI.",
         "Applies full KPI definitions (ABV, ASP, ABS, ATV, Unique Bills).",
         "Applies growth KPI rules (MoM / YoY / WoW) returning current, previous, delta, % change.",
         "Enforces brand vs category/subclass disambiguation rules (e.g. 'polo' → CATEGORY, 'USPA' → BRAND).",
         "Uses entity hints from resolved attributes with priority over LLM guessing.",
         "Output: JSON  { sql: '<T-SQL string>' }.",
     ]),
    ("4.4  SQL Debugger Agent",
     "src/agents/sql_debugger.py",
     "Repairs a failing SQL query given the database error message.",
     [
         "Triggered automatically by _execute_with_recovery() after each failed attempt.",
         "Input: original user question, failing SQL, DB error string, schema context, business context.",
         "Output: corrected SQL string (or empty string to give up).",
         "Retries are capped at sql_debug_max_retries (from settings).",
     ]),
    ("4.5  Summariser Agent",
     "src/agents/summarizer.py",
     "Converts a raw SQL result set into a concise, business-friendly natural-language answer.",
     [
         "Runs in a ThreadPoolExecutor parallel to chart generation.",
         "Input: user question, SQL used, execution result (rows + columns), business context.",
         "Produces a 2-4 sentence response suitable for a non-technical business user.",
     ]),
    ("4.6  Empty Result Handler Agent",
     "src/agents/empty_result_handler.py",
     "Explains why a SQL query returned no rows and guides the user toward a working query.",
     [
         "Triggered when execution_result.row_count == 0.",
         "Input: user question, SQL, business context, resolved entity attributes (last_entity_match).",
         "Cites specific attributes used (e.g. 'I searched using SUBCLASS = JOGGERS').",
         "Output: 3-4 sentence response — acknowledgement + probable cause + ONE corrective follow-up question.",
     ]),
    ("4.7  Chat Agent",
     "src/agents/chat.py",
     "Handles general conversational questions that do not require database queries.",
     [
         "Activated when Classifier labels message as 'normal_chat'.",
         "Input: contextual prompt with conversation history.",
         "Responds naturally without touching the database pipeline.",
     ]),
]

for title, filepath, summary, points in agents:
    add_heading(title, level=2, space_before=14)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("File: "); r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = BLUE
    r2 = p.add_run(filepath); r2.font.size = Pt(9.5); r2.font.color.rgb = GREY; r2.italic = True
    add_body(summary, space_after=3)
    for pt in points:
        add_bullet(pt)


# ══════════════════════════════════════════════════════════════════════════════
#  5. ENTITY RESOLUTION PROCESS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5.  Entity Resolution Process", level=1)
add_body(
    "Entity resolvers identify specific database values (e.g. state names, store names, product "
    "categories) mentioned in the user's question and map them to canonical column values, "
    "so the SQL Writer can filter precisely without guessing spelling."
)

add_heading("Resolution Algorithm", level=2, space_before=10)
add_step_table([
    (1, "Cache Check",         "On first call, the resolver queries the database for all distinct values in its target column and caches them (TTL-controlled)."),
    (2, "Token Extraction",    "The question is tokenised and n-grams are generated to find candidate phrases."),
    (3, "Fuzzy Match",         "SequenceMatcher compares candidate phrases against all cached values."),
    (4, "Embedding Match",     "If use_value_embeddings=True, cosine similarity between the candidate embedding and cached value embeddings is calculated."),
    (5, "Score Selection",     "The best score from fuzzy or embedding is used.  If score ≥ similarity_threshold → EntityMatch returned."),
    (6, "Hint Injection",      "All matched entities are formatted as a hint block and prepended to the LLM prompt for the SQL Writer."),
], col_widths=(0.7, 1.5, 4.1))

add_heading("Resolver Configuration", level=2, space_before=10)
add_two_col_table([
    ("StateEntityResolver",     "Column: STATE  |  Threshold: from settings  |  Embeddings: No"),
    ("CityEntityResolver",      "Column: CITY  |  Threshold: from settings  |  Embeddings: No"),
    ("StoreNameEntityResolver", "Column: STORE_NAME  |  Threshold: from settings  |  Embeddings: No"),
    ("CategoryEntityResolver",  "Column: CATEGORY  |  Threshold: min(settings, 0.80)  |  Embeddings: Yes"),
    ("SubclassEntityResolver",  "Column: SUBCLASS  |  Threshold: min(settings, 0.78)  |  Embeddings: Yes  (lower threshold tolerates misspellings e.g. 'Jogers' → JOGGERS)"),
], header=("Resolver", "Configuration"))


# ══════════════════════════════════════════════════════════════════════════════
#  6. SQL CACHE PROCESS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("6.  SQL Cache Process", level=1)
add_body(
    "The SQL cache avoids redundant and expensive LLM calls by storing previously generated "
    "SQL alongside a fingerprint of the schema and business context at the time of generation.  "
    "Two lookup strategies are used in order."
)

add_step_table([
    (1, "Question Normalisation", "normalize_question_for_cache() strips stopwords, lower-cases, and sorts tokens to produce a stable cache key regardless of minor phrasing changes."),
    (2, "Fingerprinting",         "fingerprint_text() creates an MD5/SHA hash of the current schema context and business context.  A cached SQL is only reused if both fingerprints match, ensuring cache entries are invalidated when the schema or business rules change."),
    (3, "Exact Lookup",           "find_exact() queries the SQLite store for a row matching (normalised_question, schema_fingerprint, business_fingerprint).  Returns SQL immediately if found and passes SqlGuard."),
    (4, "Semantic Lookup",        "If exact miss and sql_cache_semantic_enabled=True: the normalised question is embedded and find_semantic() computes cosine similarity against all stored embeddings with matching fingerprints.  Returns the closest match if similarity ≥ sql_cache_similarity_threshold."),
    (5, "Cache Miss → Generate",  "If both lookups miss, SQL Writer Agent is called.  The generated SQL is upserted into the cache after a successful execution."),
    (6, "Guard Before Reuse",     "Every cached SQL (exact or semantic) is passed through ensure_safe_readonly_sql() before being executed, preventing unsafe cached entries from running."),
], col_widths=(0.7, 1.8, 3.8))


# ══════════════════════════════════════════════════════════════════════════════
#  7. VISUALISATION PROCESS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("7.  Visualisation Process", level=1)
add_body(
    "Charts and tables are generated as self-contained HTML files rendered inside the Streamlit "
    "app using an iframe.  Chart type selection is data-driven and layered — explicit user "
    "requests always take priority."
)

add_step_table([
    (1, "Row Count Guard",    "If the result has fewer than 2 rows, no chart or table is generated (scalar result — answer is purely textual)."),
    (2, "Explicit Detection", "Tokenise the user question.  Check for: 'pie' → pie chart;  'line'/'trend'/'over time' → line chart;  'bar'/'chart'/'graph' → bar chart;  'table' → HTML table."),
    (3, "Semantic Pie",       "If question implies distribution ('share', 'breakdown', 'proportion') and 2 ≤ rows ≤ 12 → pie chart."),
    (4, "Semantic Line",      "If question implies trend ('monthly', 'weekly', 'daily', 'over time') and a date-like column exists → line chart."),
    (5, "Data-Driven Bar",    "If rows ≥ 2 and no explicit/semantic override matched → bar chart (default multi-row chart type)."),
    (6, "Column Detection",   "detect_date_column() and detect_label_column() identify the X-axis and grouping dimensions automatically from column names."),
    (7, "HTML Rendering",     "Chart.js HTML file is written to chart_output_dir.  For table requests, render_table_html() produces a styled HTML table with theme colours."),
    (8, "Fallback",           "On any rendering exception, a plain Python dict list row_preview is returned so the UI can display a text table."),
], col_widths=(0.7, 1.6, 4.0))

add_heading("Supported Chart Types", level=2, space_before=10)
add_two_col_table([
    ("Bar Chart",   "Best for categorical comparisons (sales by region, top stores).  Default for any multi-row result."),
    ("Line Chart",  "Best for time-series / trend data.  Requires a date/period column."),
    ("Pie Chart",   "Best for part-to-whole distributions with ≤ 12 slices."),
    ("HTML Table",  "Explicit user request ('show as table') or fallback when chart rendering fails."),
], header=("Chart Type", "When Used"))


# ══════════════════════════════════════════════════════════════════════════════
#  8. CONVERSATION MEMORY PROCESS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("8.  Conversation Memory Process", level=1)
add_body(
    "ConversationMemory provides persistent, multi-thread conversation history so the LLM "
    "has context from earlier turns, enabling follow-up questions and pronoun resolution "
    "(e.g. 'compare it with last month')."
)

add_two_col_table([
    ("Storage",          "JSON file at memory_store_path (configurable).  Persists across application restarts."),
    ("Threads",          "Each conversation session can have its own thread (identified by a UUID-based ID).  Users can create, list, and switch threads."),
    ("Turn Format",      "Each turn stores: user message, assistant reply, route (normal_chat / business_question), timestamp."),
    ("Prompt Format",    "format_for_prompt() serialises the last max_turns turns into a compact text block prepended to every LLM call."),
    ("Display Format",   "format_for_display() produces a human-readable history for the debug panel in the UI."),
    ("Auto-Create",      "If memory_auto_create_thread=True, a new thread is created automatically on startup with a timestamp-based ID."),
    ("Context Injection","_build_contextual_input() merges thread ID, history, entity hints, and the current message into a single prompt string."),
], header=("Aspect", "Details"))


# ══════════════════════════════════════════════════════════════════════════════
#  9. SQL SAFETY GUARD
# ══════════════════════════════════════════════════════════════════════════════
add_heading("9.  SQL Safety Guard", level=1)
add_body(
    "ensure_safe_readonly_sql() in src/sql_guard.py is the last line of defence before any "
    "SQL reaches the database.  It is called for both freshly generated SQL and cached SQL."
)
add_bullet("Parses the SQL string and checks the statement type.")
add_bullet("Raises SqlGuardError if any non-SELECT statement is detected (INSERT, UPDATE, DELETE, MERGE, DROP, ALTER, CREATE, TRUNCATE, EXEC, etc.).")
add_bullet("Raises SqlGuardError if any table other than prd.FACT_SALES_AI appears in FROM or JOIN clauses.")
add_bullet("SqlGuardError is caught in the orchestrator and returned to the user as a safe error message — the database is never touched.")


# ══════════════════════════════════════════════════════════════════════════════
#  10. CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("10.  Configuration Reference", level=1)
add_body("All settings are managed via src/config.py (Pydantic Settings — loaded from environment variables or .env file).")

add_two_col_table([
    ("DB_CONNECTION_STRING",           "ODBC connection string for Microsoft Fabric SQL endpoint."),
    ("LLM_MODEL",                      "LiteLLM model string for all text/JSON agent calls (e.g. azure/gpt-4o)."),
    ("EMBEDDING_MODEL",                "LiteLLM model string for embedding calls (e.g. azure/text-embedding-3-small)."),
    ("MAX_RESULT_ROWS",                "Maximum rows returned from a single SQL execution."),
    ("PREVIEW_ROWS",                   "Number of rows included in the row preview and HTML table."),
    ("SQL_CACHE_ENABLED",              "Toggle the SQL cache on/off."),
    ("SQL_CACHE_SEMANTIC_ENABLED",     "Toggle semantic (embedding) cache lookup on/off."),
    ("SQL_CACHE_SIMILARITY_THRESHOLD", "Minimum cosine similarity score for a semantic cache hit (0–1)."),
    ("SQL_DEBUG_MAX_RETRIES",          "Number of times the SQL Debugger Agent may attempt to repair a failing query."),
    ("ENTITY_SEARCH_ENABLED",          "Toggle all entity resolvers on/off."),
    ("ENTITY_STATE_SIMILARITY_THRESHOLD", "Base similarity threshold for entity resolvers."),
    ("ENTITY_STATE_CACHE_TTL_SECONDS", "How long entity value lists are cached before re-querying the database."),
    ("MEMORY_MAX_TURNS",               "Maximum conversation turns retained in the prompt context window."),
    ("VISUALIZATION_ENABLED",          "Toggle chart generation on/off."),
    ("CHART_MAX_POINTS",               "Maximum data points rendered in a single chart."),
    ("BUSINESS_CONTEXT_PATH",          "Path to the business_context.json file containing column definitions and KPI rules."),
], header=("Setting", "Description"))


# ══════════════════════════════════════════════════════════════════════════════
#  11. FILE STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("11.  Project File Structure", level=1)

file_tree = [
    ("AI-DA-AGENTS/",                        "Project root"),
    ("  app.py",                             "Streamlit UI entry point"),
    ("  business_context.json",              "Column definitions, KPI rules, brand alias mappings"),
    ("  src/",                               "Application source package"),
    ("    orchestrator.py",                  "Central ChatOrchestrator class"),
    ("    config.py",                        "Pydantic Settings — all configuration"),
    ("    models.py",                        "BotReply, SqlExecutionResult Pydantic models"),
    ("    llm.py",                           "LiteLLM wrapper: run_text_agent, run_json_agent, run_embedding"),
    ("    memory.py",                        "ConversationMemory, ThreadSummary"),
    ("    sql_cache.py",                     "SqlQueryCache, fingerprint_text, normalize_question_for_cache"),
    ("    sql_guard.py",                     "ensure_safe_readonly_sql, SqlGuardError"),
    ("    entity_search.py",                 "Five ColumnEntityResolver subclasses"),
    ("    visualization.py",                 "decide_chart, build_visual_output, render_table_html"),
    ("    business_context.py",              "BusinessContextStore, format_context_for_prompt"),
    ("    agents/",                          "Agent functions (one file per agent)"),
    ("      classifier.py",                  "classify_question()"),
    ("      clarifier.py",                   "check_needs_clarification()"),
    ("      chat.py",                        "respond_to_normal_chat()"),
    ("      sql_writer.py",                  "write_sql_query()"),
    ("      sql_debugger.py",                "debug_sql_query()"),
    ("      summarizer.py",                  "summarize_sql_result()"),
    ("      empty_result_handler.py",        "handle_empty_result()"),
    ("    db/",                              "Database access layer"),
    ("      execute.py",                     "execute_sql_query(), DatabaseExecutionError"),
    ("      schema_cache.py",               "SchemaCache — live schema retrieval and caching"),
]

add_two_col_table(file_tree, header=("Path", "Purpose"), col_widths=(2.5, 3.8))


# ══════════════════════════════════════════════════════════════════════════════
#  12. ERROR HANDLING SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("12.  Error Handling Summary", level=1)

add_two_col_table([
    ("LLM Unreachable",        "Classifier catches exception → returns 'I could not reach the LLM service' message.  All other agents fall back gracefully (clarifier → no clarification; empty handler → generic fallback string)."),
    ("SqlGuardError",          "Raised if generated or cached SQL is not a safe SELECT.  Caught in orchestrator → 'I cannot run that query safely' message.  Never reaches the database."),
    ("DatabaseExecutionError", "SQL execution failed.  Triggers SQL Debugger retry loop.  After max retries → 'I hit a database error' message returned to user."),
    ("Zero Rows Returned",     "Empty Result Handler Agent is called with the question, SQL, business context, and resolved entity attributes.  Returns a contextual 3-4 sentence follow-up."),
    ("Visualisation Failure",  "build_visual_output() catches all exceptions internally.  Falls back to a plain row_preview dict list so the UI always has something to show."),
    ("Entity Resolver Error",  "Each resolver wraps its database/embedding calls in try/except.  A failed resolver returns None (skipped) rather than crashing the pipeline."),
    ("Cache Errors",           "SQL cache read/write failures are non-fatal — the system simply falls back to generating SQL fresh via the LLM."),
], header=("Failure Scenario", "Handling Behaviour"))


# ── save ──────────────────────────────────────────────────────────────────────
out = r"C:\Users\7518549\WORK\AI-DA-AGENTS\AI_DA_AGENTS_Process_Document.docx"
doc.save(out)
print(f"Saved: {out}")
