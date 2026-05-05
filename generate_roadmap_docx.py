"""Generate SKILLS_ROADMAP.docx from the roadmap content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

# ── Color palette ─────────────────────────────────────────────────────────────
DARK_NAVY  = RGBColor(0x1A, 0x1A, 0x2E)   # titles
DARK_BLUE  = RGBColor(0x16, 0x21, 0x3E)   # H2
ACCENT     = RGBColor(0x0F, 0x3C, 0x78)   # level headers / table headers
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xFE)   # table header bg
LIGHT_GRAY = RGBColor(0xF5, 0xF7, 0xFA)   # alt row bg
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREEN      = RGBColor(0x06, 0x6D, 0x3D)
TEXT       = RGBColor(0x1F, 0x29, 0x37)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'D0D7DE')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def para(text='', bold=False, italic=False, size=11, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color
    return p

def heading1(text, emoji=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{emoji}  {text}" if emoji else text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_NAVY
    # Bottom border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0F3C78')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = ACCENT
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(f"{bold_prefix}: ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = TEXT
    return p

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    add_table_borders(table)

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, ACCENT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = LIGHT_GRAY if ri % 2 == 0 else WHITE
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = TEXT

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("🤖  Pulse AI")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = DARK_NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Developer Skills Roadmap")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = ACCENT

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A complete 9-level learning path for building a production-grade\nRAG-powered, SQL-grounded, multi-agent retail analytics chatbot")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Arvind Fashions · Internal Engineering Document · 2025")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 0 — OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

heading1("Overview", "📌")

para("Pulse AI is a production-grade conversational analytics assistant built on top of Arvind Fashions' SQL Server data warehouse. It combines large language models (LLMs), semantic SQL caching, entity resolution, multi-agent orchestration, and a React streaming UI into a single coherent system.", size=11, space_after=8)

heading2("What Pulse AI Does")
for b in [
    "Accepts natural-language questions about retail sales (brands, stores, regions, KPIs)",
    "Classifies questions as data queries vs. general chat",
    "Resolves ambiguous entity names (brand/state/city) before generating SQL",
    "Generates, executes, and verifies SQL against a live SQL Server warehouse",
    "Streams narrative answers with citations and interactive charts",
    "Maintains per-user conversation threads in MongoDB",
    "Logs every LLM call to Langfuse for cost, latency, and quality monitoring",
    "Authenticates users via Azure Active Directory / Microsoft SSO",
]:
    bullet(b)

heading2("Stack at a Glance")
make_table(
    ["Layer", "Technology"],
    [
        ["LLM Gateway",       "LiteLLM → OpenAI / Azure OpenAI"],
        ["Agent Orchestration","Python orchestrator with specialist agents"],
        ["API Layer",         "FastAPI + StreamingResponse (SSE)"],
        ["Data Warehouse",    "SQL Server via pyodbc"],
        ["Conversation Memory","MongoDB (cloud) / JSON (local)"],
        ["Semantic Cache",    "SQLite + sentence embeddings"],
        ["Entity Resolution", "difflib fuzzy + embedding similarity"],
        ["Auth",              "Azure AD · MSAL browser · python-jose JWT"],
        ["Observability",     "Langfuse v2 (self-hosted Docker)"],
        ["Frontend",          "React 19 · TypeScript · Vite · Chart.js"],
    ],
    col_widths=[2.0, 4.5]
)

heading2("How to Use This Roadmap")
for b in [
    "Work through the levels sequentially — each level builds on the previous.",
    "Time estimates assume ~10 focused hours per week; adjust to your pace.",
    "Each level ends with a 'what you can now build' summary.",
    "The Capstone Project (end of document) ties everything together.",
]:
    bullet(b)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# LEVELS
# ═══════════════════════════════════════════════════════════════════════════════

levels = [
    {
        "num": 1,
        "title": "Foundation",
        "weeks": "Weeks 1–4",
        "emoji": "🏗️",
        "outcome": "Write clean Python scripts and basic TypeScript; manage a project repo",
        "sections": [
            ("Python 3.10+ Fundamentals", [
                "Type hints: `str`, `int`, `list[T]`, `dict[K, V]`, `Optional`, `Union`",
                "Dataclasses and frozen dataclasses as lightweight config holders",
                "Generators and `yield` — essential for SSE streaming",
                "Context managers (`with` statement, `__enter__`/`__exit__`)",
                "`ContextVar` for thread-safe request-scoped state (used for trace propagation)",
                "Virtual environments: `python -m venv`, `pip install`, `requirements.txt`",
                "`python-dotenv` and `.env` file patterns",
            ]),
            ("JavaScript / TypeScript Fundamentals", [
                "TypeScript: interfaces, type aliases, generics (`Array<T>`, `Promise<T>`)",
                "Async/await and `for await...of` loops",
                "Type narrowing with `typeof`, `instanceof`, and discriminated unions",
                "ES modules: `import`/`export`, named vs. default exports",
                "npm: `package.json`, `node_modules`, scripts",
            ]),
            ("Git & Project Hygiene", [
                "Git: init, clone, add, commit, push, branch, merge",
                "`.gitignore` patterns: exclude `.env`, `__pycache__`, `node_modules`",
                "Never commit secrets — use environment variables",
                "Branching strategies: feature branches, PRs",
            ]),
        ]
    },
    {
        "num": 2,
        "title": "Backend Core (FastAPI)",
        "weeks": "Weeks 5–8",
        "emoji": "⚡",
        "outcome": "Build a REST API with streaming endpoints, validation, and CORS",
        "sections": [
            ("FastAPI", [
                "Route definitions: `@app.get`, `@app.post`, path params, query params",
                "Pydantic request/response models (automatic validation + docs)",
                "Dependency injection with `Depends()` — used for auth, settings",
                "CORS middleware: `CORSMiddleware` with allowed origins",
                "Lifespan events: `@asynccontextmanager` startup/shutdown",
                "`StreamingResponse` with `text/event-stream` content type for SSE",
                "HTTP exception handling: `HTTPException`, custom error responses",
            ]),
            ("Pydantic v2", [
                "`BaseModel`, `Field` with defaults and descriptions",
                "`model_validate()` for parsing dicts/JSON",
                "`model_dump()` for serialisation",
                "Nested models and `List[Model]` fields",
                "Validators: `@field_validator`, `@model_validator`",
            ]),
            ("Environment Configuration", [
                "`python-dotenv` loading at application startup",
                "Settings class pattern with `@lru_cache(maxsize=1)` singleton",
                "Environment-specific configs (dev vs. prod)",
            ]),
        ]
    },
    {
        "num": 3,
        "title": "Databases",
        "weeks": "Weeks 9–11",
        "emoji": "🗄️",
        "outcome": "Read from SQL Server, persist conversations in MongoDB, cache in SQLite",
        "sections": [
            ("SQL Server via pyodbc", [
                "ODBC connection strings and driver configuration",
                "Cursor operations: `execute()`, `fetchmany()`, `fetchall()`",
                "Parameterised queries to prevent SQL injection",
                "Stored procedure execution: `EXEC proc_name @param = ?`",
                "Read-only guard: reject DDL/DML statements before execution",
                "Connection pooling and error handling",
            ]),
            ("MongoDB via pymongo", [
                "Document model design: threads collection with embedded turns array",
                "Compound unique indexes: `(user_id, thread_id)`",
                "Upsert patterns: `update_one(..., upsert=True)` with `$set` / `$setOnInsert`",
                "`$push` for appending turns to arrays",
                "Per-user multi-tenancy: always filter by `user_id`",
                "Connection URI with authentication",
            ]),
            ("SQLite for Caching", [
                "Lightweight embedded database — no server required",
                "Schema design for SQL query cache (question hash → result + metadata)",
                "TTL-based invalidation: store `created_at` and compare to `time.time()`",
                "Entity search cache for fuzzy/embedding lookups",
            ]),
        ]
    },
    {
        "num": 4,
        "title": "LLM Engineering",
        "weeks": "Weeks 12–16",
        "emoji": "🧠",
        "outcome": "Write effective prompts, call LLMs via LiteLLM, handle streaming and embeddings",
        "sections": [
            ("LLM Fundamentals", [
                "System vs. user vs. assistant roles in chat completions",
                "`temperature=0` for deterministic/factual outputs; higher for creative",
                "Structured output: JSON mode, schema enforcement in system prompt",
                "Streaming completions: delta chunks, `finish_reason`",
                "Embeddings: dense vector representations, cosine similarity threshold",
                "Context window limits and token counting (`tiktoken`)",
                "Prompt injection awareness and mitigation",
            ]),
            ("LiteLLM", [
                "Multi-provider abstraction: same API for OpenAI, Azure, Anthropic, etc.",
                "`litellm.completion(model, messages, stream, response_format)`",
                "`litellm.embedding(model, input)` for vector generation",
                "`success_callback` list for observability hooks (e.g., `[\"langfuse\"]`)",
                "Metadata passing: `metadata={\"generation_name\": ..., \"trace_id\": ...}`",
                "Error handling: `litellm.exceptions.RateLimitError`, retry logic",
            ]),
            ("Agent Prompt Patterns (used in this codebase)", [
                "Classifier: binary label — business question / normal chat",
                "Domain guard: in-scope / out-of-scope check",
                "SQL writer: schema-grounded SQL generation with column list",
                "SQL explainer: plain-English explanation of query logic",
                "Summariser: streaming narrative from tabular data rows",
                "Verifier: fact-check numbers against source SQL result rows",
                "Citation builder: map each claim to a specific source row",
                "Clarifier: detect when question is too vague to answer",
                "KPI router: map question intent to a named stored procedure",
            ]),
        ]
    },
    {
        "num": 5,
        "title": "Multi-Agent Orchestration",
        "weeks": "Weeks 17–20",
        "emoji": "🎯",
        "outcome": "Build an orchestrator that routes questions through a pipeline of specialist agents",
        "sections": [
            ("Orchestrator Pattern", [
                "Central router that classifies intent and delegates to specialist agents",
                "Sequential pipeline: classify → resolve entities → generate SQL → execute → verify → summarise",
                "Parallel execution with `ThreadPoolExecutor` where agents are independent",
                "SSE streaming pipeline: SQL ready → `start` event → token stream → `metadata` event",
            ]),
            ("Error Recovery (QueryResolver)", [
                "Strategy pattern: 6 retry strategies in priority order",
                "Strategies: syntax fix, column fix, broaden filters, simplify, safe rewrite, full rewrite",
                "Each strategy passes failing SQL + error message back to LLM for correction",
                "Graceful fallbacks at every pipeline stage — never crash, always respond",
            ]),
            ("Memory Management", [
                "Conversation turns stored as a sliding window (`collections.deque` with `maxlen`)",
                "Thread isolation: create, switch, list threads per user",
                "Rich turn metadata persistence: SQL used, citations, chart data, resolver explanation",
                "Dual backend strategy: MongoDB (distributed/prod) / JSON file (local/dev)",
            ]),
            ("Semantic SQL Cache", [
                "Question normalisation: lowercase, stopword removal, strip punctuation",
                "Cache key = hash(normalised question + schema fingerprint + business context)",
                "Exact match lookup first; then embedding similarity with threshold (e.g., 0.92)",
                "Cache invalidation by TTL and manual clear endpoint",
            ]),
            ("Entity Resolution", [
                "Problem: user says 'Tommy' — resolve to 'Tommy Hilfiger' in the schema",
                "Fuzzy matching: `difflib.SequenceMatcher` ratio with confidence threshold",
                "Embedding similarity: encode entity + candidates, find nearest neighbour",
                "Resolver types: Brand, State, City, Store, Category, Subclass",
                "Confidence scoring and source tracking for explainability",
            ]),
            ("Chart / Visualisation Pipeline", [
                "Column role detection: label columns (string/date) vs. metric columns (numeric)",
                "Chart type inference from question keywords ('trend' → line, 'compare' → bar)",
                "Date format selection based on time span (daily → 'MMM D', monthly → 'MMM YY')",
                "Chart title generation from question with prefix stripping",
            ]),
        ]
    },
    {
        "num": 6,
        "title": "Authentication & Security",
        "weeks": "Weeks 21–23",
        "emoji": "🔐",
        "outcome": "Secure the API with Azure AD SSO; isolate data per user",
        "sections": [
            ("OAuth 2.0 / OpenID Connect Concepts", [
                "Authorization Code flow with PKCE (used for SPAs)",
                "Access tokens vs. ID tokens vs. refresh tokens",
                "Scopes: `openid`, `profile`, `email`, custom API scopes",
                "JWT structure: `header.payload.signature`, Base64url encoding",
                "Key JWT claims: `aud` (audience), `iss` (issuer), `exp` (expiry), `preferred_username`",
            ]),
            ("Azure Active Directory", [
                "App Registration: client ID, tenant ID, redirect URIs",
                "Platform configuration: Single-page Application (SPA) for PKCE",
                "MSAL browser v5: `PublicClientApplication`, `loginRedirect`, `acquireTokenSilent`",
                "MSAL v5 async init: `msalInstance.initialize()` BEFORE `ReactDOM.createRoot()`",
                "ID token as Bearer for backend validation",
            ]),
            ("Backend JWT Validation (python-jose)", [
                "Fetch Microsoft JWKS from `https://login.microsoftonline.com/{tenant}/discovery/v2.0/keys`",
                "In-memory JWKS cache with 1-hour TTL (avoid fetching on every request)",
                "RS256 signature verification: `jwt.decode(token, jwks, algorithms=['RS256'], audience=..., issuer=...)`",
                "Extract `preferred_username` → use as `user_id` throughout the system",
                "Graceful dev-mode bypass when `AZURE_TENANT_ID` is unset",
            ]),
            ("Per-User Isolation", [
                "Per-user orchestrator registry: `dict[str, ChatOrchestrator]` + `threading.Lock`",
                "Double-checked locking pattern for thread-safe lazy initialisation",
                "Per-user MongoDB namespace: all queries filter on `user_id`",
                "Per-user thread list: each user sees only their own conversations",
            ]),
        ]
    },
    {
        "num": 7,
        "title": "Frontend (React + TypeScript)",
        "weeks": "Weeks 24–27",
        "emoji": "🖥️",
        "outcome": "Build a streaming chat UI with auth, charts, and thread management",
        "sections": [
            ("React 19", [
                "Function components with hooks: `useState`, `useEffect`, `useRef`, `useCallback`",
                "Async data loading patterns in `useEffect`",
                "Token-by-token streaming UI: append to message text on each SSE token",
                "Component composition: App → Sidebar + Main → Message + ChatInput",
                "Refs: `useRef` for DOM access (scroll-to-bottom) and interval cleanup",
            ]),
            ("TypeScript in React", [
                "Interface definitions for API response types (`ChatResponse`, `Thread`, `Citation`)",
                "Generic types for event handlers and state: `useState<MessageItem[]>`",
                "Type narrowing in event handlers",
                "Union types for message roles: `'user' | 'bot'`",
                "Type-only imports: `import type { Foo }` vs. `import { Foo }`",
            ]),
            ("MSAL React", [
                "`MsalProvider` wrapping the root — passes MSAL context to all children",
                "`AuthenticatedTemplate` / `UnauthenticatedTemplate` conditional rendering",
                "`useMsal()` hook for access to MSAL instance and accounts",
                "`acquireTokenSilent` → falls back to `loginRedirect` on error",
                "ID token extraction: `resp.idToken` sent as `Authorization: Bearer ...`",
            ]),
            ("SSE Client (fetch-based)", [
                "`fetch()` with `ReadableStream` body reader — no EventSource (can't set headers)",
                "`TextDecoder` for converting `Uint8Array` chunks to strings",
                "Line-by-line parsing: split on `\\n`, handle `data: ` prefix",
                "Async generator pattern: `async function* streamMessage(text)` yields typed events",
                "Typed events: `start` | `token` | `metadata` | `complete` | `error`",
            ]),
            ("Chart.js + react-chartjs-2", [
                "Chart types: `Line`, `Bar`, `Pie` from react-chartjs-2",
                "Dataset configuration: labels, data arrays, backgroundColor, borderColor",
                "Responsive sizing with `maintainAspectRatio: false`",
                "Dynamic chart type selection based on `chart_type` from API response",
            ]),
            ("Vite", [
                "Dev server with Hot Module Replacement (HMR)",
                "Proxy configuration in `vite.config.ts`: `/api` → backend URL",
                "`VITE_*` environment variables injected at build time",
                "TypeScript compilation and bundling",
            ]),
        ]
    },
    {
        "num": 8,
        "title": "Observability & Feedback (Langfuse)",
        "weeks": "Week 28",
        "emoji": "📊",
        "outcome": "Track every LLM call, visualise costs and latency, collect user feedback",
        "sections": [
            ("Langfuse v2 SDK", [
                "Trace lifecycle: `lf.trace(name, user_id, input)` → spans (auto via LiteLLM) → `lf.trace(id, output)` → `lf.flush()`",
                "LiteLLM auto-logging: `litellm.success_callback = ['langfuse']` — zero boilerplate",
                "`ContextVar` trace propagation: set `trace_id` once at request start, read in every LLM call",
                "`metadata` dict passed to LiteLLM: `{'trace_id': ..., 'generation_name': ..., 'user_id': ...}`",
                "User feedback: `lf.score(trace_id=..., name='user-feedback', value=1.0, data_type='BOOLEAN')`",
                "Graceful degradation: all Langfuse calls wrapped in `if _lf:` checks",
            ]),
            ("Self-Hosting with Docker", [
                "Langfuse v2 image: `langfuse/langfuse:2` (NOT `:latest` which is v3+)",
                "v2 only requires Postgres — simple `docker-compose.yml` with 2 services",
                "v3 requires ClickHouse + Redis + MinIO — significantly more complex",
                "Dashboard on port 3001; create project → copy public + secret keys",
                "Environment variables: `LANGFUSE_PUBLIC_KEY`, `LANGFUSE_SECRET_KEY`, `LANGFUSE_HOST`",
            ]),
            ("Monitoring Concepts", [
                "Per-user token spend: filter Langfuse dashboard by `user_id`",
                "Per-agent latency: each LiteLLM call becomes a named span",
                "Cost breakdown by model and agent type",
                "Feedback loop: `score` objects linked to exact trace → measure answer quality over time",
            ]),
        ]
    },
    {
        "num": 9,
        "title": "DevOps & Infrastructure",
        "weeks": "Week 29",
        "emoji": "🚀",
        "outcome": "Package, deploy, and operate the full stack in a containerised environment",
        "sections": [
            ("Docker & Docker Compose", [
                "Dockerfile for Python backend: multi-stage build, non-root user",
                "Dockerfile for React frontend: Vite build → Nginx static serve",
                "Docker Compose: define services, networks, volumes, health checks",
                "Health checks: `test: ['CMD', 'curl', '-f', 'http://localhost:8000/health']`",
                "Restart policies: `restart: unless-stopped` for production services",
                "Image pinning: always pin to a specific tag (`:2`, `:16-alpine`) not `:latest`",
            ]),
            ("Environment Management", [
                "`.env` for local secrets — NEVER commit to git",
                "`.env.example` with placeholder values — commit this as documentation",
                "Docker Compose `env_file` directive to load `.env` into containers",
                "`SSL_CERT_FILE` environment variable for corporate proxy / TLS inspection",
                "Secret management patterns: Azure Key Vault, AWS Secrets Manager",
            ]),
        ]
    },
]

for level in levels:
    n = level["num"]
    heading1(f"Level {n} — {level['title']}  ({level['weeks']})", level["emoji"])

    # Outcome box
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("🎯 After this level you can: ")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREEN
    r2 = p.add_run(level["outcome"])
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GREEN

    for section_title, bullets in level["sections"]:
        heading2(section_title)
        for b in bullets:
            bullet(b)

    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CAPSTONE PROJECT
# ═══════════════════════════════════════════════════════════════════════════════

heading1("Capstone Project — Build a Mini Pulse AI", "🏆")

para("Apply every skill from the roadmap by building a smaller but production-worthy version of Pulse AI end-to-end. Work through these sprints sequentially:", size=11, space_after=8)

make_table(
    ["Sprint", "Goal", "Key Skills"],
    [
        ["1 — API skeleton",     "FastAPI with /chat endpoint, health check, CORS",     "Level 2"],
        ["2 — SQL integration",  "Connect to SQLite/SQL Server, execute agent-written SQL", "Levels 3, 4"],
        ["3 — LLM agent",        "Question → SQL → plain-English answer via LiteLLM",   "Level 4"],
        ["4 — SSE streaming",    "Stream tokens to client; progress events",             "Levels 2, 7"],
        ["5 — React frontend",   "Chat input, message list, streaming token display",    "Level 7"],
        ["6 — Memory",           "Last 5 turns stored in MongoDB or JSON file",          "Level 5"],
        ["7 — Auth",             "Azure AD login, JWT validation, per-user isolation",   "Level 6"],
        ["8 — Observability",    "Langfuse traces, cost tracking, thumbs up/down",       "Level 8"],
        ["9 — Docker",           "docker-compose.yml for backend + frontend + DB",       "Level 9"],
    ],
    col_widths=[1.6, 2.8, 1.6]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# LIBRARY CHEAT SHEET
# ═══════════════════════════════════════════════════════════════════════════════

heading1("Quick-Reference Library Cheat Sheet", "📚")

heading2("Python Libraries")
make_table(
    ["Package", "Version (approx)", "Purpose"],
    [
        ["fastapi",           "≥ 0.111",    "HTTP API framework"],
        ["uvicorn",           "≥ 0.29",     "ASGI server for FastAPI"],
        ["pydantic",          "v2",          "Data validation and settings"],
        ["python-dotenv",     "≥ 1.0",      "Load .env files"],
        ["litellm",           "≥ 1.40",     "Multi-provider LLM gateway"],
        ["pyodbc",            "≥ 5.1",      "SQL Server connectivity"],
        ["pymongo",           "≥ 4.7",      "MongoDB driver"],
        ["python-jose[cryptography]", "≥ 3.3", "JWT decode and verification"],
        ["httpx",             "≥ 0.27",     "Async HTTP client (JWKS fetch)"],
        ["langfuse",          "≥ 2, < 3",   "LLM observability (v2 SDK)"],
        ["sentence-transformers", "≥ 2.6",  "Local embedding models"],
        ["python-docx",       "≥ 1.1",      "Generate Word .docx files"],
    ],
    col_widths=[2.2, 1.4, 3.0]
)

heading2("JavaScript / TypeScript Packages")
make_table(
    ["Package", "Version (approx)", "Purpose"],
    [
        ["react",                "19",          "UI component framework"],
        ["typescript",           "≥ 5",         "Typed JavaScript"],
        ["vite",                 "≥ 5",         "Build tool and dev server"],
        ["@azure/msal-browser",  "≥ 3 (v5 API)","MSAL browser auth"],
        ["@azure/msal-react",    "≥ 2",         "React hooks for MSAL"],
        ["axios",                "≥ 1.7",       "HTTP client"],
        ["chart.js",             "≥ 4",         "Chart rendering"],
        ["react-chartjs-2",      "≥ 5",         "React wrapper for Chart.js"],
        ["uuid",                 "≥ 9",         "Generate unique message IDs"],
        ["react-markdown",       "≥ 9",         "Render Markdown in chat bubbles"],
    ],
    col_widths=[2.2, 1.4, 3.0]
)

heading2("Infrastructure")
make_table(
    ["Tool", "Purpose"],
    [
        ["Docker Desktop",        "Local container runtime"],
        ["Docker Compose",        "Multi-container orchestration"],
        ["MongoDB Atlas / local", "Conversation persistence"],
        ["SQL Server",            "Production data warehouse"],
        ["Langfuse (self-hosted)","LLM observability dashboard"],
        ["Azure Active Directory","Enterprise SSO and JWT issuance"],
        ["Nginx",                 "Static file server for React build"],
    ],
    col_widths=[2.2, 4.4]
)

# ── Footer ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Pulse AI · Arvind Fashions Engineering · Generated 2025")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
r.italic = True

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\7518549\WORK\AI-DA-AGENTS\SKILLS_ROADMAP.docx"
doc.save(output_path)
print(f"OK  Saved: {output_path}")
